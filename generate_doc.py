import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from datetime import datetime

# Logo (asegúrate de que esté en 'assets/logo_mgi.png')
LOGO_PATH = os.path.join("assets", "logo_mgi.png")

def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_header_with_logo(doc, logo_path=LOGO_PATH):
    if not os.path.exists(logo_path):
        print("⚠️ Logo no encontrado")
        return False

    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.5))
    return True

def create_word_doc(results, name):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)

    add_header_with_logo(doc)

    now = datetime.now()
    doc.add_heading("Informe de Debida Diligencia", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Este informe contiene capturas de pantalla de:")
    doc.add_paragraph("• Lista de sanciones OFAC")
    doc.add_paragraph("• Búsquedas temáticas en Bing (español e inglés)")
    doc.add_paragraph(f"Fecha: {now.strftime('%d/%m/%Y')} — Hora: {now.strftime('%H:%M')}")

    add_page_break(doc)
    doc.add_heading(name, level=1)

    # Agrupar por tipo
    ofac = [r for r in results if r['type'] == 'OFAC']
    bing = [r for r in results if r['type'] == 'Bing']

    # OFAC section
    if ofac:
        r = ofac[0]
        doc.add_heading("Búsqueda en lista OFAC", level=2)
        if r["path"] and os.path.exists(r["path"]):
            doc.add_picture(r["path"], width=Inches(5.0))
            msg = "⚠️ Coincidencias encontradas." if r.get("has_results") else "✓ Sin coincidencias."
            doc.add_paragraph(msg)
        else:
            doc.add_paragraph("❌ No se pudo capturar la búsqueda en OFAC.")
        add_page_break(doc)

    # Bing section
    doc.add_heading("Búsquedas en Bing por categoría de riesgo", level=2)

    for lang in ["es", "en"]:
        doc.add_heading("Español" if lang == "es" else "English", level=3)
        results_lang = [r for r in bing if r["lang"] == lang]

        for r in results_lang:
            doc.add_heading(f"Categoría: {r['category']}", level=4)

            # Mostrar términos
            doc.add_paragraph("Términos buscados:", style='List Bullet')
            doc.add_paragraph(r['criterios'])

            # Mostrar imagen
            if r["path"] and os.path.exists(r["path"]):
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(r["path"], width=Inches(4.75))
            else:
                doc.add_paragraph("❌ Imagen no disponible.")

            add_page_break(doc)

    # Exportar a memoria
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
